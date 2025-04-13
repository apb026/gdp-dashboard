import streamlit as st
import google.generativeai as genai
from io import BytesIO
from docx import Document
import requests
from PIL import Image
import io
import base64
import time
from transformers import pipeline


# Function to generate code documentation using Gemini API
def generate_code_documentation(code_input, api_key):
    try:
        code_content = f"Here's a code snippet: {code_input} \n\n---\n\n Can you generate documentation for this code?"

        # Generate documentation using the Gemini API
        response_code = genai.Client(api_key=api_key).models.generate_content(
            model="gemini-2.0-flash", 
            contents=[{"parts": [{"text": code_content}]}]
        )

        if response_code.candidates:
            doc_answer = response_code.candidates[0].content.parts[0].text
            return doc_answer
        return None
    except Exception as e:
        st.error(f"Error generating code documentation: {str(e)}")
        return None

# # Commented above function and incorporated a new function which uses a hugging face model for chat completion.
# def generate_code_documentation(code_input):
#     try:
#         # Initialize a text generation pipeline with a code generation model
#         generator = pipeline("text-generation", model="facebook/incoder-1B", tokenizer="facebook/incoder-1B")
        
#         # Prepare the input text for generating documentation
#         code_content = f"Here's a code snippet: {code_input} \n\n---\n\n Can you generate documentation for this code?"
        
#         # Generate documentation
#         doc_answer = generator(code_content, max_length=512)[0]['generated_text']
#         return doc_answer
#     except Exception as e:
#         st.error(f"Error generating code documentation: {str(e)}")
#         return None

# Function to add styled text in the Word document
def add_styled_text(doc, text, style=None, is_bold=False, is_italic=False, is_code=False):
    para = doc.add_paragraph()
    
    if is_code:
        run = para.add_run(text)
        run.font.name = 'Courier New'  # Use monospace font for code
        para.style = 'Normal'
    else:
        if style:
            para.style = style
        run = para.add_run(text)
        
        # Apply bold or italic formatting if needed
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True

# Function to generate document answer with few-shot prompting
def generate_document_answer_with_few_shot(document_text, question):
    examples = """
    Example 1:
    Document: "Python is a high-level programming language that is easy to learn and use."
    Question: "What is Python?"
    Answer: "Python is a high-level programming language known for its simplicity and readability."

    Example 2:
    Document: "Machine learning is a subset of artificial intelligence where computers learn from data."
    Question: "What is machine learning?"
    Answer: "Machine learning is a branch of artificial intelligence where algorithms use data to improve their performance over time."
    """

    persona = "You are a helpful assistant trained to provide detailed, well-structured answers based on the content of the document."
    content = f"{persona} Below are a few examples of how I answer questions based on document content:\n{examples}\n\nDocument: {document_text}\nQuestion: {question}\nAnswer:"

    return content

# # Function to generate project/report file (e.g., Class 12th Chemistry lab experiments)
# def generate_project_report(subject, api_key):
#     try:
#         # Create a new document for the project report
#         doc = Document()
#         doc.add_heading(f"{subject} Project Report", 0)

#         # Use Gemini API to generate project details (Lab experiments, for example)
#         prompt = f"Generate a list of detailed lab experiments for Class 12th Chemistry, including the experiment details and images where possible."
#         response = genai.Client(api_key=api_key).models.generate_content(
#             model="gemini-2.0-flash", 
#             contents=[{"parts": [{"text": prompt}]}]
#         )

#         if response.candidates:
#             # Extract the generated text
#             project_content = response.candidates[0].content.parts[0].text
#             add_styled_text(doc, project_content, style="Normal")

#             # Generate images or diagrams for each experiment (use the description for generating images)
#             experiment_descriptions = project_content.split("\n")  # Assuming each experiment description is on a new line
#             for description in experiment_descriptions:
#                 # Generate image for the experiment using Gemini (or other APIs like DALL·E or any text-to-image API)
#                 image_url = generate_image_from_prompt(description, api_key)
#                 if image_url:
#                     add_styled_text(doc, f"Experiment: {description}", style="Heading 2", is_bold=True)
#                     insert_image_to_word(doc, image_url)
        
#             # Save the project report as a Word document
#             doc_io = BytesIO()
#             doc.save(doc_io)
#             doc_io.seek(0)

#             # Provide download button
#             # st.download_button(
#             #     label="Download Project Report with Experiments and Images",
#             #     data=doc_io,
#             #     file_name="chemistry_project_report.docx",
#             #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#             # )
#         else:
#             st.error("Failed to generate the project content.")

#     except Exception as e:
#         st.error(f"An error occurred while generating the project report: {str(e)}")

# Using third model from huggingface here
def generate_project_report(subject):
    try:
        # Initialize a text generation pipeline
        generator = pipeline("text-generation", model="EleutherAI/gpt-neo-2.7B", tokenizer="EleutherAI/gpt-neo-2.7B")
        
        # Prepare the prompt for generating a report
        prompt = f"Generate a detailed report on the subject: {subject}. Include relevant experiments, steps, and explanations."
        
        # Generate the report
        project_content = generator(prompt, max_length=512)[0]['generated_text']
        
        return project_content
    except Exception as e:
        st.error(f"Error generating project report: {str(e)}")
        return None


# Function to generate an image based on the description (Gemini or other APIs)
def generate_image_from_prompt(prompt, api_key, retries=3):
    try:
        response = genai.Client(api_key=api_key).models.generate_content(
            model="gemini-2.0-flash", 
            contents=[{"parts": [{"text": f"Generate an image based on the following description: {prompt}"}]}]
        )
        if response.candidates:
            image_data = response.candidates[0].content.parts[0].text
            return image_data
        else:
            st.error("Failed to generate image.")
            return None
    except Exception as e:
        if retries > 0 and "RESOURCE_EXHAUSTED" in str(e):
            st.warning("API quota exceeded. Retrying...")
            time.sleep(5)  # wait for 5 seconds before retrying
            return generate_image_from_prompt(prompt, api_key, retries - 1)
        else:
            st.error(f"Error generating image: {str(e)}")
            return None


# Function to insert image into Word document
def insert_image_to_word(doc, image_url):
    try:
        # If it's a base64 string, decode and save as image
        if image_url.startswith('data:image'):
            image_data = base64.b64decode(image_url.split(",")[1])
            image = Image.open(BytesIO(image_data))
        else:
            image = Image.open(requests.get(image_url, stream=True).raw)
        
        # Save image to a BytesIO buffer to insert into doc
        img_stream = BytesIO()
        image.save(img_stream, format="PNG")
        img_stream.seek(0)
        
        doc.add_picture(img_stream)
    except Exception as e:
        st.error(f"Error inserting image into document: {str(e)}")

# Streamlit UI
st.title("📄 Document Question Answering, Code Documentation")

# User input for the API key
gemini_api_key = st.text_input("Gemini API Key", type="password")
if not gemini_api_key:
    st.info("Please add your Gemini API key to continue.", icon="🗝️")
else:
    # Create a Gemini client
    client = genai.Client(api_key=gemini_api_key)

    user_choice = st.radio(
        "Select the functionality you want to use:",
        ("Upload Document for Q&A", "Provide Code for Documentation")
    )

    if user_choice == "Upload Document for Q&A":
        # Handle the document upload and question
        uploaded_file = st.file_uploader(
            "Upload a document (.txt, .md, .pdf, .docx)", type=("txt", "md", "pdf", "docx")
        )
        
        question = st.text_area(
            "Now ask a question about the document!",
            placeholder="Can you give me a short summary?",
            disabled=not uploaded_file,
        )

        if uploaded_file and question:
            try:
                document = uploaded_file.read()

                # Handle different file types (PDF, DOCX)
                if uploaded_file.type == "application/pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(BytesIO(document))
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    document = text

                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    from docx import Document
                    doc = Document(BytesIO(document))
                    text = ""
                    for para in doc.paragraphs:
                        text += para.text
                    document = text

                # Generate content with few-shot prompt for document Q&A
                content = generate_document_answer_with_few_shot(document, question)

                # Generate an answer using the Gemini API for the document Q&A
                response = client.models.generate_content(
                    model="gemini-2.0-flash", 
                    contents=[{"parts": [{"text": content}]}]
                )

                # Access the first candidate and its text
                if response.candidates:
                    answer = response.candidates[0].content.parts[0].text
                    st.write(answer)
                else:
                    st.error("No response from the model.")
                
            except Exception as e:
                st.error(f"An error occurred while processing the document: {str(e)}")

    elif user_choice == "Provide Code for Documentation":
        # Handle the code input for generating documentation
        code_input = st.text_area(
            "Paste your code here to generate documentation.",
            placeholder="def add(a, b):\n    return a + b",
            height=200
        )
        
        generate_code_doc = st.button("Generate Documentation for Code")

        if code_input and generate_code_doc:
            try:
                # Generate code documentation
                doc_answer = generate_code_documentation(code_input, gemini_api_key)
                # doc_answer = generate_code_documentation(code_input)
                
                if doc_answer:
                    st.write(doc_answer)
                    
                    # Create a Word document with the generated documentation
                    doc = Document()
                    doc.add_heading('Code Documentation', 0)

                    # Add introductory sections
                    add_styled_text(doc, "Overview", style="Heading 1", is_bold=True)
                    add_styled_text(doc, "This document provides a detailed explanation of the code snippet provided by the user.", is_italic=True)
                    doc.add_paragraph("\n")

                    # Parse the generated documentation and add formatted text
                    add_styled_text(doc, doc_answer)

                    # Save the document in memory
                    doc_io = BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)

                    # Provide download link for the Word document
                    st.download_button(
                        label="Download Documentation as Word File",
                        data=doc_io,
                        file_name="code_documentation.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                else:
                    st.error("No response from the model.")
                
            except Exception as e:
                st.error(f"An error occurred while generating code documentation: {str(e)}")

    # elif user_choice == "Generate Cover Image":
    #     # Get the cover image description from the user
    #     cover_image_description = st.text_area(
    #         "Describe the cover image you want to generate.",
    #         placeholder="A Golden Retriever playing on a sunny beach.",
    #         height=200
    #     )
        
        # generate_cover_image = st.button("Generate Cover Image")

        # if cover_image_description and generate_cover_image:
        #     try:
        #         # Generate cover image based on the description
        #         image_url = generate_image_from_prompt(cover_image_description, gemini_api_key)

        #         if image_url:
        #             st.image(image_url, caption="Generated Cover Image", use_container_width=True)
        #         else:
        #             st.error("No image was generated. Please provide a more detailed description.")
                
        #     except Exception as e:
        #         st.error(f"An error occurred while generating the cover image: {str(e)}")